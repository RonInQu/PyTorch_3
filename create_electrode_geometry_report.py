"""
Generate a .docx report summarizing the Gemini conversation about
Electrode Geometry for Impedance Measurement in a PA catheter.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Electrode Geometry for Impedance Measurement', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Summary of Gemini AI Research Session (June 20, 2026)\n'
    'Application: Endovascular catheter for pulmonary artery embolism detection\n'
    'Tissue classification: Blood / Clot / Vessel Wall via bioimpedance',
    style='Subtitle'
)

# ============================================================
# 1. SURFACE TOPOGRAPHY
# ============================================================
doc.add_heading('1. Surface Topography: Dimpled vs. Flat Electrodes', level=1)

doc.add_heading('Dimpled Electrodes', level=2)
bullets = [
    'Increased Effective Surface Area: Dimples increase microscopic surface area without enlarging '
    'the electrode footprint. Per the Helmholtz model, Z_contact ∝ 1/A — larger area lowers '
    'electrode-tissue interface impedance.',
    'Mechanical Anchoring (Dry): Dimples/micro-pillars pierce through high-impedance layers, '
    'establishing more stable electrical connections.',
    'Reduced Motion Artifacts: Textured surface "locks" into tissue, reducing lateral sliding.',
    'Boundary Layer Disruption (intravascular): Micro-turbulence disrupts the hydrodynamic '
    'boundary layer, ensuring fresh blood contacts the electrode.',
    'CRITICAL DRAWBACK — Thrombogenicity: Dimples create micro-stagnation zones where platelets '
    'aggregate, risking catheter-induced thrombosis.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Flat Electrodes', level=2)
bullets = [
    'Uniform Current Density: Highly predictable, uniform distribution assuming ideal contact.',
    'Ease of Cleaning & Reusability: Flat surfaces are easier to sanitize.',
    'Hemodynamic Neutrality: Flush-mounted electrodes maintain laminar flow, preventing '
    'localized clotting and shear-stress artifacts.',
    'Predictable E-Field Penetration: Clean, well-mapped fringe field lines simplify '
    'inverse problem algorithms for tissue classification.',
    'Susceptibility to Blood "Shunting": If not pressed firmly, surrounding blood acts as '
    'electrical short circuit, masking underlying tissue signature.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 2. CURVATURE
# ============================================================
doc.add_heading('2. Curvature: Concave vs. Convex', level=1)

doc.add_heading('Convex Electrodes (Domed / Protruding)', level=2)
bullets = [
    'Wall & Clot Indentation: Physically pushes past ambient blood film, forcing direct '
    'contact with solid tissue matrix.',
    'High-Density Focus: Current density spikes at the apex, providing highly localized '
    '"probe" behavior with excellent spatial resolution.',
    'Ideal for discriminating clot boundary from true lumen.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Concave Electrodes (Recessed Bands / Wells)', level=2)
bullets = [
    'Blood-Only Reference Sensors: Shielded from direct wall/clot contact.',
    'Differential Sensing Strategy: Continuously calibrate ambient blood impedance baseline, '
    'enabling forward-facing electrodes to isolate the delta caused by a thrombus.',
    'Ideal as shaft-mounted reference electrodes.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 3. GEOMETRIC SHAPE
# ============================================================
doc.add_heading('3. Geometric Shape: Circular vs. Rectangular', level=1)

doc.add_heading('Circular Electrodes', level=2)
bullets = [
    'Symmetrical Electric Field: Radially symmetric — eliminates orientation bias. '
    'Impedance reading invariant to catheter rotation.',
    'Minimal Edge Effects: No sharp corners eliminates current crowding hot-spots.',
    'Mathematical Modeling: Analytical solutions (e.g., Newman\'s formula) '
    'are simplest for circular disk geometries.',
    'Isotropic Signature: Contact force and angle against soft PE won\'t distort '
    'the phase angle signature θ.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Rectangular (Ring or Strip) Electrodes', level=2)
bullets = [
    'Axi-Symmetrical Sensing (Bands): 360° radial sensing field — detects '
    'wall/clot regardless of catheter rotation.',
    'Axial Tracking (Strips): Can map the length of a thrombus as device passes through.',
    'High Edge-Effect Crowding: Sharp corners cause localized heating and polarization '
    'errors, especially at frequencies <1 kHz.',
    'Space Efficiency: Dense arrays without unmeasured gaps.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 4. SUMMARY TABLE
# ============================================================
doc.add_heading('4. Selection Guide Summary', level=1)

table = doc.add_table(rows=7, cols=3)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Feature', 'Best For', 'Main Drawback']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ('Dimpled', 'Maximizing surface area, disrupting boundary layer',
     'Thrombogenicity risk; difficult to clean'),
    ('Flat', 'Hemodynamic neutrality, clean field penetration, easy cleaning',
     'Blood shunting if not in firm contact'),
    ('Concave', 'Blood reference sensors, differential baseline calibration',
     'Cannot contact target tissue directly'),
    ('Convex', 'Direct tissue contact, displacing blood film, spatial resolution',
     'Current crowding at apex; spacing-dependent shunting'),
    ('Circular', 'Isotropic measurement, rotation-invariant, minimal edge effects',
     'Poor space utilization in tight linear arrays'),
    ('Rectangular', 'Directional current, axial mapping, 360° ring sensing',
     'Corner current crowding; sensitive to rotational alignment'),
]
for i, (feat, best, drawback) in enumerate(data, start=1):
    table.rows[i].cells[0].text = feat
    table.rows[i].cells[1].text = best
    table.rows[i].cells[2].text = drawback

# ============================================================
# 5. BLOOD FILM / SHUNT EFFECT
# ============================================================
doc.add_heading('5. The Highly Conductive Blood Film ("Blood Shunt")', level=1)

doc.add_paragraph(
    'When an electrode approaches a clot or vessel wall, a microscopic film of blood '
    '(σ ≈ 0.7 S/m) remains trapped at the interface. Current preferentially passes '
    'laterally through this highly conductive film rather than penetrating into the '
    'more resistive clot or wall tissue.'
)

doc.add_heading('Interface States and Signal Interpretation', level=2)
tbl2 = doc.add_table(rows=5, cols=4)
tbl2.style = 'Medium Shading 1 Accent 1'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
h2 = ['Interface State', '|Z| Magnitude', 'Phase Angle (high freq)', 'Interpretation']
for i, h in enumerate(h2):
    tbl2.rows[0].cells[i].text = h
states = [
    ('Pure Blood (No Contact)', 'Extremely Low', 'Minimal (resistive)', 'Open Lumen'),
    ('Partial Contact (Thick Film)', 'Low to Medium', 'Low phase shift', 'Transition Zone'),
    ('Firm Contact on Clot', 'High', 'High negative phase shift', 'Thrombus Detected'),
    ('Firm Contact on Wall', 'Medium-High', 'Moderate phase shift', 'Vessel Wall Reached'),
]
for i, (state, mag, phase, interp) in enumerate(states, start=1):
    tbl2.rows[i].cells[0].text = state
    tbl2.rows[i].cells[1].text = mag
    tbl2.rows[i].cells[2].text = phase
    tbl2.rows[i].cells[3].text = interp

doc.add_paragraph(
    'Key indicator: A sudden transition from purely resistive low impedance to a complex, '
    'capacitive signature indicates the electrode has broken through the blood film and '
    'is reading the actual embolic mass.'
)

# ============================================================
# 6. ELECTRICAL DOUBLE LAYER (EDL)
# ============================================================
doc.add_heading('6. Electrical Double Layer (EDL) & Electrode Polarization', level=1)

doc.add_heading('EDL Structure', level=2)
doc.add_paragraph(
    'When a metallic electrode contacts blood, charge rearrangement creates the EDL:'
)
edl_items = [
    'Electrode Surface: Net electronic charge.',
    'Inner Helmholtz Plane (IHP): Specifically adsorbed ions (dehydrated anions) '
    'on the metal surface.',
    'Outer Helmholtz Plane (OHP): Solvated hydrated counter-ions (Na⁺, K⁺) at '
    '~angstrom distance from metal.',
    'Diffuse Layer (Gouy-Chapman): Ion concentration gradient decaying to bulk.'
]
for item in edl_items:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Equivalent Circuit Model', level=2)
doc.add_paragraph(
    'The total measured impedance:\n\n'
    '    Z_total = 2 · Z_CPE + Z_blood/clot\n\n'
    'Where the Constant Phase Element (CPE) models the non-ideal double-layer:\n\n'
    '    Z_CPE = 1 / [Q₀ · (jω)^α]\n\n'
    '• Q₀: pseudo-capacitance related to C_dl (typically 10–40 μF/cm²)\n'
    '• α: homogeneity factor (1.0 = ideal capacitor; ~0.8 for textured surfaces)\n'
    '• ω: angular frequency (2πf)'
)

doc.add_heading('Frequency Impact on EDL', level=2)
freq_items = [
    'Low Frequencies (<10 kHz) — EDL Dominated: Z_CPE is astronomically high. '
    '90%+ of voltage drops across the sub-nm EDL, not the tissue. '
    'Cannot reliably differentiate clot from blood.',
    'Intermediate to High (10 kHz–1 MHz) — Bypassing EDL: Z_CPE → 0. '
    'EDL acts as short-circuit, allowing current to penetrate bulk tissue. '
    'Ideal window for tissue classification.',
]
for item in freq_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Geometry × EDL Interactions', level=2)
geo_edl = [
    'Dimpled electrodes lower EDL impedance: Increased area A → increased Q₀ → '
    'polarization corner frequency shifts downward → clean data at lower frequencies.',
    'Rectangular/Sharp edges distort EDL: Uneven current density causes non-uniform '
    'charging, lowering α and complicating signal deconvolution.',
    'Protein Fouling: Fibrinogen/albumin/platelets adsorb onto electrode (IHP/OHP), '
    'altering Q₀ over time. Multi-frequency spectroscopy required to dynamically '
    'isolate and subtract shifting Z_CPE baseline.',
]
for item in geo_edl:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 7. DUAL CIRCULAR CONVEX CONFIGURATION
# ============================================================
doc.add_heading('7. Dual Circular Convex Configuration (±1.5V, 50 kHz)', level=1)

doc.add_heading('Why Dual-Circular is Mathematically Optimal', level=2)
bullets = [
    'Symmetrical Current Density: No sharp corners → no hot-spots.',
    'Elimination of Orientation Bias: Invariant to catheter rotation in PA branches.',
    'Simplified Modeling: Analytical solutions readily available for parallel circular disks.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Why Dual-Convex is Physically Optimal', level=2)
bullets = [
    'Both domes compress tissue, squeeze out blood film, minimize shunting simultaneously.',
    'High current density at apex provides localized probing.',
    'CATCH: If electrodes too close, current takes shortest path through blood gap '
    'between apices without penetrating into tissue.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Electrochemical Safety at 50 kHz', level=2)
doc.add_paragraph(
    'Water electrolysis threshold: 1.23V. At 3.0V differential, electrolysis risk exists '
    'if DC or low-frequency AC is used. At 50 kHz:'
)
safety_items = [
    'Voltage alternates every 20 μs (10 μs per half-cycle).',
    'Ions drift only fractions of a nanometer before reversing → no net Faradaic '
    'electron transfer → no gas bubble formation.',
    '50 kHz is "safe but borderline" — the system is protected from electrolysis, '
    'but the EDL is not fully shorted out (some residual polarization artifact remains).',
    'Operating at 200 kHz–1 MHz would further reduce EDL contribution.',
]
for item in safety_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('50 kHz Assessment Summary', level=2)
doc.add_paragraph(
    'At 50 kHz with ±1.5V on dual circular convex electrodes:\n'
    '• Electrochemistry: SAFE — no electrolysis or bubble formation\n'
    '• EDL bypass: PARTIAL — ~70-80% of signal comes from bulk tissue; '
    'residual polarization artifact still present\n'
    '• β-dispersion sensing: YES — 50 kHz sits within the β-dispersion band '
    '(10 kHz–10 MHz) where cell membrane capacitance produces the structural '
    'signatures that differentiate clot from wall\n'
    '• Blood shunt mitigation: Convex geometry provides mechanical displacement; '
    'sufficient spacing ensures field penetration depth'
)

# ============================================================
# 8. RECOMMENDED ARCHITECTURE
# ============================================================
doc.add_heading('8. Recommended Design Architecture for PE Detection', level=1)

doc.add_paragraph(
    'Mixed-geometry approach for highest sensitivity:'
)

rec_items = [
    'Active Sensing (Tip): Slightly convex, circular micro-electrodes on the '
    'active face. Convexity displaces blood film during contact; circular shape '
    'focuses current without corner-polarization artifacts.',
    'Blood Calibration (Shaft): Flush or slightly concave rectangular ring '
    'electrodes along shaft to capture baseline ambient blood impedance, '
    'enabling real-time differential subtraction.',
    'Spacing: Sufficient electrode separation to force field lines deep into '
    'tissue rather than jumping across shortest blood gap.',
    'Drive: ±1.5V high-frequency AC sine wave with zero DC bias.',
]
for i, item in enumerate(rec_items, start=1):
    doc.add_paragraph(f'{i}. {item}')

# ============================================================
# SAVE
# ============================================================
output_dir = r'c:\Users\RonaldKurnik\OneDrive - Inquis Medical\Documents\2026\PyTorch_3\Comsol'
output_path = os.path.join(output_dir, 'Electrode_Geometry_Impedance_Report.docx')
doc.save(output_path)
print(f'Report saved: {output_path}')
